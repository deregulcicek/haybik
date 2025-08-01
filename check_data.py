from docx import Document
import pandas as pd

def check_data_structure():
    """Train ve test verilerinin yapısını kontrol eder"""
    
    # Train verisi
    print("=== TRAIN VERİSİ ANALİZİ ===")
    doc_train = Document('interviews_train.docx')
    
    train_texts = []
    for para in doc_train.paragraphs:
        if para.text.strip():
            train_texts.append(para.text.strip())
    
    print(f"Train verisi paragraf sayısı: {len(train_texts)}")
    print("İlk 5 paragraf:")
    for i, text in enumerate(train_texts[:5]):
        print(f"{i+1}. {text[:100]}...")
    
    # Test verisi
    print("\n=== TEST VERİSİ ANALİZİ ===")
    doc_test = Document('interviews_test.docx')
    
    test_texts = []
    for para in doc_test.paragraphs:
        if para.text.strip():
            test_texts.append(para.text.strip())
    
    print(f"Test verisi paragraf sayısı: {len(test_texts)}")
    print("İlk 5 paragraf:")
    for i, text in enumerate(test_texts[:5]):
        print(f"{i+1}. {text[:100]}...")
    
    # Toplam analiz
    total_samples = len(train_texts) + len(test_texts)
    print(f"\n=== TOPLAM ANALİZ ===")
    print(f"Toplam örnek sayısı: {total_samples}")
    print(f"Train örnek sayısı: {len(train_texts)}")
    print(f"Test örnek sayısı: {len(test_texts)}")
    
    # Katılımcı hesaplama (9 soru per katılımcı varsayımı)
    if total_samples % 9 == 0:
        participants = total_samples // 9
        print(f"Tahmini katılımcı sayısı: {participants}")
    else:
        print("9 soru per katılımcı varsayımı doğrulanamadı")

if __name__ == "__main__":
    check_data_structure() 