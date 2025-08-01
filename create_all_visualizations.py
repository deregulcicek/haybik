import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier, GradientBoostingClassifier, VotingClassifier
from sklearn.svm import SVC
from sklearn.metrics import roc_auc_score, precision_recall_fscore_support, accuracy_score
from sklearn.preprocessing import LabelEncoder
from imblearn.over_sampling import SMOTE
import re
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from docx import Document
import warnings
warnings.filterwarnings('ignore')

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

# Set style for publication-quality plots
plt.style.use('seaborn-v0_8-whitegrid')
sns.set_palette("husl")

def load_interview_data():
    """Load interview data from DOCX files"""
    # Load train data
    doc_train = Document('interviews_train.docx')
    train_texts = []
    for para in doc_train.paragraphs:
        text = para.text.strip()
        if text and not text.startswith('Interview') and not text.startswith('Interviewer:'):
            if 'Interviewee:' in text:
                response = text.split('Interviewee:')[-1].strip()
                if response:
                    train_texts.append(response)
            elif len(text) > 20:
                train_texts.append(text)
    
    # Load test data
    doc_test = Document('interviews_test.docx')
    test_texts = []
    for para in doc_test.paragraphs:
        text = para.text.strip()
        if text and not text.startswith('Interview') and not text.startswith('Interviewer:'):
            if 'Interviewee:' in text:
                response = text.split('Interviewee:')[-1].strip()
                if response:
                    test_texts.append(response)
            elif len(text) > 20:
                test_texts.append(text)
    
    return train_texts, test_texts

def enhanced_text_preprocessing(text):
    """Enhanced text preprocessing with comprehensive noise reduction"""
    if pd.isna(text) or text == '':
        return ''
    
    text = str(text).lower()
    
    # Fix contractions
    contractions = {
        "i'm": "i am", "you're": "you are", "he's": "he is", "she's": "she is",
        "it's": "it is", "we're": "we are", "they're": "they are",
        "i've": "i have", "you've": "you have", "we've": "we have",
        "they've": "they have", "i'll": "i will", "you'll": "you will",
        "he'll": "he will", "she'll": "she will", "it'll": "it will",
        "we'll": "we will", "they'll": "they will", "i'd": "i would",
        "you'd": "you would", "he'd": "he would", "she'd": "she would",
        "it'd": "it would", "we'd": "we would", "they'd": "they would",
        "isn't": "is not", "aren't": "are not", "wasn't": "was not",
        "weren't": "were not", "hasn't": "has not", "haven't": "have not",
        "hadn't": "had not", "doesn't": "does not", "don't": "do not",
        "didn't": "did not", "won't": "will not", "wouldn't": "would not",
        "shan't": "shall not", "shouldn't": "should not", "can't": "cannot",
        "couldn't": "could not", "mustn't": "must not"
    }
    
    for contraction, expansion in contractions.items():
        text = text.replace(contraction, expansion)
    
    text = re.sub(r'[^\w\s]', ' ', text)
    tokens = word_tokenize(text)
    stop_words = set(stopwords.words('english'))
    
    extended_noise_words = {
        'im', 'makes', 'checking', 'get', 'go', 'going', 'got', 'gets',
        'make', 'made', 'making', 'check', 'checks', 'checked',
        'thing', 'things', 'way', 'ways', 'time', 'times', 'day', 'days',
        'people', 'person', 'someone', 'somebody', 'anyone', 'anybody',
        'everyone', 'everybody', 'nothing', 'something', 'anything',
        'everything', 'here', 'there', 'where', 'when', 'why', 'how',
        'what', 'which', 'who', 'whom', 'whose', 'this', 'that', 'these',
        'those', 'yes', 'no', 'maybe', 'okay', 'ok', 'right', 'wrong',
        'good', 'bad', 'nice', 'great', 'awesome', 'cool', 'fine',
        'sure', 'definitely', 'absolutely', 'probably', 'possibly',
        'really', 'very', 'quite', 'rather', 'pretty', 'so', 'too',
        'also', 'just', 'only', 'even', 'still', 'already', 'yet',
        'again', 'back', 'forward', 'up', 'down', 'in', 'out', 'on',
        'off', 'over', 'under', 'above', 'below', 'between', 'among',
        'through', 'across', 'around', 'about', 'against', 'toward',
        'towards', 'into', 'onto', 'upon', 'within', 'without',
        'before', 'after', 'during', 'since', 'until', 'while',
        'because', 'although', 'though', 'unless', 'if', 'else',
        'then', 'than', 'as', 'like', 'such', 'same', 'different',
        'other', 'another', 'each', 'every', 'all', 'both', 'either',
        'neither', 'some', 'any', 'many', 'much', 'few', 'several',
        'most', 'least', 'more', 'less', 'most', 'least', 'first',
        'second', 'third', 'last', 'next', 'previous', 'current',
        'new', 'old', 'young', 'big', 'small', 'large', 'little',
        'high', 'low', 'long', 'short', 'wide', 'narrow', 'deep',
        'shallow', 'thick', 'thin', 'heavy', 'light', 'strong',
        'weak', 'hard', 'soft', 'easy', 'difficult', 'simple',
        'complex', 'clear', 'confusing', 'obvious', 'hidden',
        'visible', 'invisible', 'open', 'closed', 'full', 'empty',
        'busy', 'free', 'available', 'unavailable', 'possible',
        'impossible', 'necessary', 'unnecessary', 'important',
        'unimportant', 'useful', 'useless', 'helpful', 'harmful',
        'safe', 'dangerous', 'clean', 'dirty', 'fresh', 'stale',
        'hot', 'cold', 'warm', 'cool', 'dry', 'wet', 'smooth',
        'rough', 'quiet', 'loud', 'bright', 'dark', 'light',
        'heavy', 'fast', 'slow', 'quick', 'rapid', 'gradual',
        'sudden', 'gradual', 'constant', 'variable', 'regular',
        'irregular', 'normal', 'abnormal', 'usual', 'unusual',
        'common', 'rare', 'frequent', 'infrequent', 'often',
        'sometimes', 'rarely', 'never', 'always', 'usually',
        'generally', 'typically', 'normally', 'usually', 'often',
        'sometimes', 'occasionally', 'rarely', 'seldom', 'hardly',
        'scarcely', 'barely', 'almost', 'nearly', 'about', 'around',
        'approximately', 'roughly', 'exactly', 'precisely', 'just',
        'only', 'merely', 'simply', 'purely', 'entirely', 'completely',
        'totally', 'fully', 'partially', 'mostly', 'mainly', 'primarily',
        'chiefly', 'largely', 'considerably', 'significantly', 'substantially',
        'noticeably', 'obviously', 'clearly', 'evidently', 'apparently',
        'seemingly', 'supposedly', 'allegedly', 'reportedly', 'apparently',
        'obviously', 'clearly', 'evidently', 'apparently', 'seemingly',
        'supposedly', "allegedly", "reportedly", "apparently", "obviously",
        "clearly", "evidently", "apparently", "seemingly", "supposedly",
        "allegedly", "reportedly", "apparently", "obviously", "clearly",
        "evidently", "apparently", "seemingly", "supposedly", "allegedly",
        "reportedly", "apparently", "obviously", "clearly", "evidently",
        "apparently", "seemingly", "supposedly", "allegedly", "reportedly"
    }
    
    all_stop_words = stop_words.union(extended_noise_words)
    
    filtered_tokens = []
    for token in tokens:
        if (len(token) >= 3 and 
            token.lower() not in all_stop_words and 
            not token.isdigit()):
            filtered_tokens.append(token.lower())
    
    return ' '.join(filtered_tokens)

def create_synthetic_dataset():
    """Create synthetic dataset with enhanced diversity"""
    
    type1_negative = [
        "I feel like the algorithm is controlling what I see",
        "The algorithm seems to manipulate my feed content",
        "I notice the algorithm steering me toward certain posts",
        "The algorithm appears to control my social media experience",
        "I feel the algorithm is influencing my content choices",
        "The algorithm seems to decide what I should see",
        "I notice algorithmic control over my feed",
        "The algorithm appears to manipulate my preferences",
        "I feel like the algorithm is programming my choices",
        "The algorithm seems to control my social media usage",
        "I'm being steered by the algorithm without realizing it",
        "The algorithm is shaping my online experience",
        "I feel trapped by algorithmic decisions",
        "The algorithm determines what content I consume",
        "I'm being guided by invisible algorithmic forces"
    ]
    
    type2_negative = [
        "I feel disconnected from real conversations",
        "Social media makes me feel isolated sometimes",
        "I'm losing touch with genuine human connections",
        "The digital world feels separate from real life",
        "I feel alienated from authentic social experiences",
        "Social media creates artificial social bonds",
        "I'm becoming disconnected from reality",
        "The digital space feels different from actual life",
        "I feel like I'm living in a virtual world",
        "Real human interaction feels different now",
        "I'm more comfortable with digital interactions than real ones",
        "The online world feels more real than offline life",
        "I prefer digital communication over face-to-face",
        "I feel more connected to my phone than to people",
        "The digital realm has replaced my social life"
    ]
    
    type3_negative = [
        "I check social media more often than I should",
        "I feel dependent on these digital platforms",
        "Social media has become my main source of information",
        "I'm constantly checking my phone for updates",
        "I feel like I need social media to stay connected",
        "My daily routine revolves around social media",
        "I'm trapped in a cycle of platform usage",
        "I can't imagine life without these platforms",
        "Social media controls my daily activities",
        "I feel addicted to checking social media",
        "I can't go more than an hour without checking my phone",
        "My mood depends on social media engagement",
        "I feel anxious when I can't access social media",
        "I spend more time online than with real people",
        "Social media has become my primary source of validation"
    ]
    
    type4_negative = [
        "I only see content that matches my beliefs",
        "The algorithm creates an echo chamber effect",
        "I'm stuck in a filter bubble of similar opinions",
        "I rarely encounter challenging perspectives",
        "The algorithm reinforces my existing views",
        "I'm living in an ideological echo chamber",
        "I only see one side of most issues",
        "The algorithm prevents diverse viewpoints",
        "I'm trapped in a confirmation bias loop",
        "I rarely get exposed to opposing views",
        "My feed only shows content I already agree with",
        "I never see dissenting opinions anymore",
        "The algorithm keeps me in my comfort zone",
        "I'm surrounded by people who think like me",
        "I miss hearing different perspectives"
    ]
    
    positive_expressions = [
        "I appreciate how the algorithm helps me discover content",
        "The algorithm seems to understand my preferences",
        "I find the personalized recommendations useful",
        "The algorithm makes my experience better",
        "I like how the algorithm curates content",
        "The algorithm helps me find relevant information",
        "I enjoy the personalized content suggestions",
        "The algorithm enhances my social media usage",
        "I find the algorithmic recommendations helpful",
        "The algorithm improves my overall experience",
        "The algorithm saves me time by showing relevant content",
        "I appreciate the personalized experience",
        "The algorithm helps me discover new interests",
        "I find the content curation helpful",
        "The algorithm makes my social media more enjoyable"
    ]
    
    neutral_expressions = [
        "I use social media to stay connected with friends",
        "I check social media occasionally throughout the day",
        "I post content when I have something to share",
        "I follow accounts that interest me",
        "I use social media for entertainment purposes",
        "I scroll through my feed when I'm bored",
        "I share photos and updates with my network",
        "I use social media to keep up with current events",
        "I interact with content that appears in my feed",
        "I use social media as a communication tool",
        "I use social media to stay updated with friends",
        "I occasionally post updates about my life",
        "I use social media to find interesting content",
        "I follow various accounts for different reasons",
        "I use social media as part of my daily routine"
    ]
    
    synthetic_data = []
    
    for text in type1_negative + type2_negative + type3_negative + type4_negative:
        synthetic_data.append({'text': text, 'sentiment': 'negative', 'ssa_type': 'type1-4'})
    
    for text in positive_expressions:
        synthetic_data.append({'text': text, 'sentiment': 'positive', 'ssa_type': 'positive'})
    
    for text in neutral_expressions:
        synthetic_data.append({'text': text, 'sentiment': 'neutral', 'ssa_type': 'neutral'})
    
    return pd.DataFrame(synthetic_data)

def run_analysis():
    """Run the complete analysis to get results"""
    
    # Load original data
    train_texts, test_texts = load_interview_data()
    
    # Create synthetic dataset
    synthetic_df = create_synthetic_dataset()
    
    # Combine original train data with synthetic data
    original_train_df = pd.DataFrame({
        'text': train_texts,
        'sentiment': 'neutral',
        'ssa_type': 'original'
    })
    
    combined_train_df = pd.concat([original_train_df, synthetic_df], ignore_index=True)
    
    # Create test dataset
    original_test_df = pd.DataFrame({
        'text': test_texts,
        'sentiment': 'neutral',
        'ssa_type': 'original'
    })
    
    test_synthetic = synthetic_df.sample(n=30, random_state=42)
    test_df = pd.concat([original_test_df, test_synthetic], ignore_index=True)
    
    # Preprocess text
    combined_train_df['clean_text'] = combined_train_df['text'].apply(enhanced_text_preprocessing)
    test_df['clean_text'] = test_df['text'].apply(enhanced_text_preprocessing)
    
    # Prepare features and labels
    X_train = combined_train_df['clean_text']
    y_train = combined_train_df['sentiment']
    X_test = test_df['clean_text']
    y_test = test_df['sentiment']
    
    # Encode labels
    label_encoder = LabelEncoder()
    y_train_encoded = label_encoder.fit_transform(y_train)
    y_test_encoded = label_encoder.transform(y_test)
    
    # Feature engineering
    tfidf_vectorizer = TfidfVectorizer(
        max_features=300,
        ngram_range=(1, 2),
        min_df=3,
        max_df=0.8,
        stop_words='english'
    )
    
    count_vectorizer = CountVectorizer(
        max_features=200,
        ngram_range=(1, 2),
        min_df=3,
        max_df=0.8,
        stop_words='english'
    )
    
    X_train_tfidf = tfidf_vectorizer.fit_transform(X_train)
    X_test_tfidf = tfidf_vectorizer.transform(X_test)
    
    X_train_count = count_vectorizer.fit_transform(X_train)
    X_test_count = count_vectorizer.transform(X_test)
    
    from scipy.sparse import hstack
    X_train_combined = hstack([X_train_tfidf, X_train_count])
    X_test_combined = hstack([X_test_tfidf, X_test_count])
    
    # Apply SMOTE
    smote = SMOTE(k_neighbors=5, random_state=42)
    X_train_balanced, y_train_balanced = smote.fit_resample(X_train_combined, y_train_encoded)
    
    # Train models
    models = {
        'Logistic Regression': LogisticRegression(random_state=42, max_iter=1500),
        'Random Forest': RandomForestClassifier(random_state=42),
        'Gradient Boosting': GradientBoostingClassifier(random_state=42),
        'SVM': SVC(random_state=42, probability=True)
    }
    
    results = {}
    
    for name, model in models.items():
        model.fit(X_train_balanced, y_train_balanced)
        
        y_pred = model.predict(X_test_combined)
        y_pred_proba = model.predict_proba(X_test_combined)
        
        y_pred_labels = label_encoder.inverse_transform(y_pred)
        
        accuracy = accuracy_score(y_test, y_pred_labels)
        precision, recall, f1, _ = precision_recall_fscore_support(y_test, y_pred_labels, average='weighted')
        
        try:
            roc_auc = roc_auc_score(y_test_encoded, y_pred_proba, multi_class='ovr')
        except:
            roc_auc = 0.0
        
        results[name] = {
            'model': model,
            'accuracy': accuracy,
            'precision': precision,
            'recall': recall,
            'f1': f1,
            'roc_auc': roc_auc,
            'y_pred': y_pred_labels,
            'y_pred_proba': y_pred_proba
        }
    
    return results, tfidf_vectorizer, combined_train_df, test_df, X_test_combined, y_test

def create_sentiment_distribution_plot(train_df):
    """Create sentiment distribution pie chart"""
    plt.figure(figsize=(10, 8))
    
    sentiment_counts = train_df['sentiment'].value_counts()
    colors = ['#ff6b6b', '#4ecdc4', '#45b7d1']
    
    plt.pie(sentiment_counts.values, labels=sentiment_counts.index, 
            autopct='%1.1f%%', colors=colors, startangle=90,
            textprops={'fontsize': 12, 'fontweight': 'bold'})
    
    plt.title('Sentiment Distribution in Training Dataset', 
              fontsize=16, fontweight='bold', pad=20)
    
    plt.tight_layout()
    plt.savefig('sentiment_distribution.png', dpi=300, bbox_inches='tight')
    plt.close()
    print("✅ Sentiment distribution plot saved as 'sentiment_distribution.png'")

def create_model_performance_plot(results):
    """Create model performance comparison bar chart"""
    plt.figure(figsize=(12, 8))
    
    model_names = list(results.keys())
    accuracies = [results[name]['accuracy'] for name in model_names]
    f1_scores = [results[name]['f1'] for name in model_names]
    
    x = np.arange(len(model_names))
    width = 0.35
    
    bars1 = plt.bar(x - width/2, accuracies, width, label='Accuracy', 
                    color='#ff6b6b', alpha=0.8)
    bars2 = plt.bar(x + width/2, f1_scores, width, label='F1-Score', 
                    color='#4ecdc4', alpha=0.8)
    
    plt.xlabel('Machine Learning Models', fontsize=12, fontweight='bold')
    plt.ylabel('Performance Score', fontsize=12, fontweight='bold')
    plt.title('Model Performance Comparison', fontsize=16, fontweight='bold')
    plt.xticks(x, model_names, rotation=45)
    plt.legend(fontsize=11)
    plt.grid(True, alpha=0.3)
    
    # Add value labels on bars
    for bars in [bars1, bars2]:
        for bar in bars:
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
                    f'{height:.3f}', ha='center', va='bottom', fontsize=10)
    
    plt.tight_layout()
    plt.savefig('model_performance.png', dpi=300, bbox_inches='tight')
    plt.close()
    print("✅ Model performance plot saved as 'model_performance.png'")

def create_roc_curves_plot(results, y_test):
    """Create ROC curves for each class"""
    plt.figure(figsize=(10, 8))
    
    # Get best model (SVM)
    best_model_name = 'SVM'
    best_model_results = results[best_model_name]
    y_pred_proba = best_model_results['y_pred_proba']
    
    from sklearn.preprocessing import label_binarize
    from sklearn.metrics import roc_curve, auc
    
    classes = np.unique(y_test)
    y_test_bin = label_binarize(y_test, classes=classes)
    
    colors = ['#ff6b6b', '#4ecdc4', '#45b7d1']
    
    for i, color in enumerate(colors):
        if i < len(classes):
            fpr, tpr, _ = roc_curve(y_test_bin[:, i], y_pred_proba[:, i])
            roc_auc = auc(fpr, tpr)
            plt.plot(fpr, tpr, color=color, lw=2,
                    label=f'{classes[i]} (AUC = {roc_auc:.3f})')
    
    plt.plot([0, 1], [0, 1], 'k--', lw=2)
    plt.xlim([0.0, 1.0])
    plt.ylim([0.0, 1.05])
    plt.xlabel('False Positive Rate', fontsize=12, fontweight='bold')
    plt.ylabel('True Positive Rate', fontsize=12, fontweight='bold')
    plt.title('ROC Curves by Sentiment Class', fontsize=16, fontweight='bold')
    plt.legend(loc="lower right", fontsize=11)
    plt.grid(True, alpha=0.3)
    
    plt.tight_layout()
    plt.savefig('roc_curves.png', dpi=300, bbox_inches='tight')
    plt.close()
    print("✅ ROC curves plot saved as 'roc_curves.png'")

def create_feature_importance_plot(results, vectorizer):
    """Create feature importance analysis"""
    plt.figure(figsize=(12, 8))
    
    # Get best model (SVM doesn't have feature_importances_, so use Random Forest)
    best_model = results['Random Forest']['model']
    feature_importance = best_model.feature_importances_
    
    # Get feature names (only TF-IDF features for visualization)
    feature_names = vectorizer.get_feature_names_out()
    feature_importance = feature_importance[:len(feature_names)]
    
    # Get top features
    max_features = min(15, len(feature_names))
    top_indices = np.argsort(feature_importance)[-max_features:]
    top_features = feature_names[top_indices]
    top_importance = feature_importance[top_indices]
    
    plt.barh(range(len(top_features)), top_importance, color='#45b7d1', alpha=0.8)
    plt.yticks(range(len(top_features)), top_features, fontsize=11)
    plt.xlabel('Feature Importance Score', fontsize=12, fontweight='bold')
    plt.title('Top 15 Most Important Features for SSA Classification', 
              fontsize=16, fontweight='bold')
    plt.grid(True, alpha=0.3)
    
    plt.tight_layout()
    plt.savefig('feature_importance.png', dpi=300, bbox_inches='tight')
    plt.close()
    print("✅ Feature importance plot saved as 'feature_importance.png'")

def create_comprehensive_dashboard(results, vectorizer, train_df, test_df, X_test_vec, y_test):
    """Create comprehensive analysis dashboard"""
    fig, axes = plt.subplots(2, 2, figsize=(16, 12))
    fig.suptitle('SSA Analysis Dashboard', fontsize=18, fontweight='bold')
    
    # 1. Sentiment Distribution
    ax1 = axes[0, 0]
    sentiment_counts = train_df['sentiment'].value_counts()
    colors = ['#ff6b6b', '#4ecdc4', '#45b7d1']
    ax1.pie(sentiment_counts.values, labels=sentiment_counts.index, 
            autopct='%1.1f%%', colors=colors, startangle=90)
    ax1.set_title('Sentiment Distribution', fontweight='bold', fontsize=14)
    
    # 2. Model Performance
    ax2 = axes[0, 1]
    model_names = list(results.keys())
    accuracies = [results[name]['accuracy'] for name in model_names]
    f1_scores = [results[name]['f1'] for name in model_names]
    
    x = np.arange(len(model_names))
    width = 0.35
    
    bars1 = ax2.bar(x - width/2, accuracies, width, label='Accuracy', 
                    color='#ff6b6b', alpha=0.8)
    bars2 = ax2.bar(x + width/2, f1_scores, width, label='F1-Score', 
                    color='#4ecdc4', alpha=0.8)
    
    ax2.set_xlabel('Models')
    ax2.set_ylabel('Score')
    ax2.set_title('Model Performance Comparison', fontweight='bold', fontsize=14)
    ax2.set_xticks(x)
    ax2.set_xticklabels(model_names, rotation=45)
    ax2.legend()
    ax2.grid(True, alpha=0.3)
    
    # 3. Feature Importance
    ax3 = axes[1, 0]
    best_model = results['Random Forest']['model']
    feature_importance = best_model.feature_importances_
    feature_names = vectorizer.get_feature_names_out()
    feature_importance = feature_importance[:len(feature_names)]
    
    max_features = min(10, len(feature_names))
    top_indices = np.argsort(feature_importance)[-max_features:]
    top_features = feature_names[top_indices]
    top_importance = feature_importance[top_indices]
    
    ax3.barh(range(len(top_features)), top_importance, color='#45b7d1', alpha=0.8)
    ax3.set_yticks(range(len(top_features)))
    ax3.set_yticklabels(top_features)
    ax3.set_xlabel('Feature Importance')
    ax3.set_title('Top 10 Feature Importance', fontweight='bold', fontsize=14)
    ax3.grid(True, alpha=0.3)
    
    # 4. ROC Curves
    ax4 = axes[1, 1]
    best_model_results = results['SVM']
    y_pred_proba = best_model_results['y_pred_proba']
    
    from sklearn.preprocessing import label_binarize
    from sklearn.metrics import roc_curve, auc
    
    classes = np.unique(y_test)
    y_test_bin = label_binarize(y_test, classes=classes)
    
    colors = ['#ff6b6b', '#4ecdc4', '#45b7d1']
    for i, color in enumerate(colors):
        if i < len(classes):
            fpr, tpr, _ = roc_curve(y_test_bin[:, i], y_pred_proba[:, i])
            roc_auc = auc(fpr, tpr)
            ax4.plot(fpr, tpr, color=color, lw=2,
                    label=f'{classes[i]} (AUC = {roc_auc:.3f})')
    
    ax4.plot([0, 1], [0, 1], 'k--', lw=2)
    ax4.set_xlim([0.0, 1.0])
    ax4.set_ylim([0.0, 1.05])
    ax4.set_xlabel('False Positive Rate')
    ax4.set_ylabel('True Positive Rate')
    ax4.set_title('ROC Curves by Class', fontweight='bold', fontsize=14)
    ax4.legend(loc="lower right")
    ax4.grid(True, alpha=0.3)
    
    plt.tight_layout()
    plt.savefig('ssa_analysis_dashboard.png', dpi=300, bbox_inches='tight')
    plt.close()
    print("✅ Comprehensive dashboard saved as 'ssa_analysis_dashboard.png'")

def main():
    """Main function to create all visualizations"""
    print("🚀 Starting comprehensive SSA analysis and visualization...")
    
    # Run analysis
    results, vectorizer, train_df, test_df, X_test_vec, y_test = run_analysis()
    
    print("\n📊 Creating individual visualizations...")
    
    # Create individual plots
    create_sentiment_distribution_plot(train_df)
    create_model_performance_plot(results)
    create_roc_curves_plot(results, y_test)
    create_feature_importance_plot(results, vectorizer)
    
    print("\n📈 Creating comprehensive dashboard...")
    create_comprehensive_dashboard(results, vectorizer, train_df, test_df, X_test_vec, y_test)
    
    print("\n✅ All visualizations completed successfully!")
    print("\n📁 Generated files:")
    print("   - sentiment_distribution.png")
    print("   - model_performance.png") 
    print("   - roc_curves.png")
    print("   - feature_importance.png")
    print("   - ssa_analysis_dashboard.png")

if __name__ == "__main__":
    main() 